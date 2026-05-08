from docx import Document

def export_docx(paras: list, out_path: str) -> None:
    """Write paragraph dicts to a .docx file at out_path."""
    doc = Document()
    for p in paras:
        para      = doc.add_paragraph(style=p["style"])
        run       = para.add_run(p["text"])
        run.bold  = p["bold"]
        para.alignment = p["align"]
    doc.save(out_path)
